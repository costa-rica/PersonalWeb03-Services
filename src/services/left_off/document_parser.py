"""
Document parser for extracting content from LEFT-OFF.docx file.
"""

import logging
import re
from datetime import datetime, timedelta
from docx import Document


logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for extracting last 7 days of activities from LEFT-OFF.docx."""

    def __init__(self, docx_path):
        """
        Initialize document parser.

        Args:
            docx_path: Path to the .docx file
        """
        self.docx_path = docx_path
        self.document = None

    def load_document(self):
        """
        Load the .docx document.

        Returns:
            bool: True if successful, False otherwise
        """
        try:
            logger.info(f"Loading document: {self.docx_path}")
            self.document = Document(self.docx_path)
            logger.info(f"Document loaded successfully ({len(self.document.paragraphs)} paragraphs)")
            return True
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return False

    def extract_last_7_days(self, output_path):
        """
        Extracts content from the last 7 days and saves to a markdown file.

        The document has Heading 1 sections in YYYYMMDD format. This method finds
        the first Heading 1 that is 8+ days old and extracts all content above it.

        Args:
            output_path: Path to save the extracted markdown content

        Returns:
            bool: True if successful, False otherwise
        """
        if not self.document:
            logger.error("Document not loaded")
            return False

        # Calculate the cutoff date (8 days ago)
        cutoff_date = datetime.now() - timedelta(days=8)
        cutoff_date_str = cutoff_date.strftime('%Y%m%d')
        logger.info(f"Extracting content from last 7 days (cutoff: {cutoff_date_str})")

        # Find the cutoff paragraph index
        cutoff_index = None
        date_pattern = re.compile(r'^\d{8}$')  # YYYYMMDD format

        for i, paragraph in enumerate(self.document.paragraphs):
            # Check if this is a Heading 1
            if paragraph.style.name == 'Heading 1':
                text = paragraph.text.strip()

                # Check if it matches YYYYMMDD format
                if date_pattern.match(text):
                    logger.debug(f"Found Heading 1 date: {text}")

                    # Check if this date is 8+ days ago
                    if text <= cutoff_date_str:
                        cutoff_index = i
                        logger.info(f"Found cutoff date: {text} at paragraph {i}")
                        break

        if cutoff_index is None:
            logger.warning("No cutoff date found - extracting entire document")
            cutoff_index = len(self.document.paragraphs)

        # Extract content from beginning to cutoff
        logger.info(f"Extracting {cutoff_index} paragraphs")
        content_lines = []

        for i in range(cutoff_index):
            paragraph = self.document.paragraphs[i]
            text = paragraph.text

            # Add markdown formatting based on style
            if paragraph.style.name == 'Heading 1':
                content_lines.append(f"# {text}")
            elif paragraph.style.name == 'Heading 2':
                content_lines.append(f"## {text}")
            elif paragraph.style.name == 'Heading 3':
                content_lines.append(f"### {text}")
            else:
                content_lines.append(text)

        # Join with newlines and save
        content = '\n'.join(content_lines)

        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            logger.info(f"Extracted content saved to: {output_path}")
            logger.info(f"Content length: {len(content)} characters")
            return True
        except Exception as e:
            logger.error(f"Failed to save extracted content: {e}")
            return False
